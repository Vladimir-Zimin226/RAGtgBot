from langchain_community.document_loaders import PyPDFLoader, CSVLoader
from langchain.schema import Document as LangDocument
from docx import Document as DocxReader
import pandas as pd
import os


def load_document(file_path: str):
    """
    Загрузка документа разных форматов в список LangDocument.
    Поддерживаются PDF, CSV, DOCX, XLSX и XLS.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return PyPDFLoader(file_path).load()
    elif ext == '.csv':
        return CSVLoader(file_path).load()
    elif ext in ('.xlsx', '.xls'):
        engine = 'openpyxl' if ext == '.xlsx' else 'xlrd'
        sheets = pd.read_excel(file_path, sheet_name=None, engine=engine)
        docs = []
        for sheet_name, df in sheets.items():
            rows = df.fillna('').astype(str).values.tolist()
            text_lines = ['\t'.join(row) for row in rows]
            text = f"Sheet: {sheet_name}\n" + '\n'.join(text_lines)
            docs.append(
                LangDocument(
                    page_content=text,
                    metadata={"source": os.path.basename(file_path), "sheet": sheet_name}
                )
            )
        return docs
    elif ext == '.docx':
        doc = DocxReader(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text]
        text = '\n'.join(full_text)
        return [LangDocument(page_content=text, metadata={"source": os.path.basename(file_path)})]
    else:
        raise ValueError(f"Не поддерживаемый формат файла: {ext}")